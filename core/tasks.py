import os
from celery import shared_task
from django.conf import settings
from docx import Document

from .models import FileUpload, ActivityLog


@shared_task
def process_file_task(file_id):
    """
    Reads an uploaded file (.txt or .docx), counts words,
    updates the FileUpload model, and logs the activity.
    """
    try:
        file_obj = FileUpload.objects.get(id=file_id)

        file_path = file_obj.file.path
        extension = os.path.splitext(file_path)[1].lower()

        word_count = 0
        if extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                word_count = len(text.split())

        elif extension == ".docx":
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            word_count = len(text.split())

        file_obj.word_count = word_count
        file_obj.status = "completed"
        file_obj.save()

        # Log activity
        ActivityLog.objects.create(
            user=file_obj.user,
            action="file_processed",
            metadata={"file_id": file_obj.id, "word_count": word_count}
        )

    except Exception as e:
        # Mark file as failed
        FileUpload.objects.filter(id=file_id).update(status="failed")

        ActivityLog.objects.create(
            user=file_obj.user if 'file_obj' in locals() else None,
            action="file_processing_failed",
            metadata={"error": str(e)}
        )
        raise e


def process_file_wordcount(file_id):
    """
    Counts words in the uploaded file and updates the FileUpload model.
    """
    try:
        file_obj = FileUpload.objects.get(id=file_id)

        file_path = file_obj.file.path
        extension = os.path.splitext(file_path)[1].lower()

        word_count = 0
        if extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                word_count = len(text.split())

        elif extension == ".docx":
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            word_count = len(text.split())

        file_obj.word_count = word_count
        file_obj.save()

        # Log activity
        ActivityLog.objects.create(
            user=file_obj.user,
            action="file_wordcounted",
            metadata={"file_id": file_obj.id, "word_count": word_count}
        )

    except Exception as e:
        # Mark file as failed
        FileUpload.objects.filter(id=file_id).update(status="failed")

        ActivityLog.objects.create(
            user=file_obj.user if 'file_obj' in locals() else None,
            action="file_wordcounting_failed",
            metadata={"error": str(e)}
        )
        raise e



